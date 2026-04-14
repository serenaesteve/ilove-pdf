import os
import io
import zipfile
import tempfile
from flask import Flask, render_template, request, send_file, jsonify
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image
import img2pdf

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

ALLOWED_PDF = {'pdf'}
ALLOWED_IMG = {'jpg', 'jpeg', 'png', 'webp'}
ALLOWED_DOCX = {'docx'}

def allowed(filename, exts):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in exts

TOOLS = {
    'merge':      {'title': 'Unir PDF',       'desc': 'Combina múltiples PDFs en uno solo',          'icon': '⊕', 'accept': '.pdf',                  'multiple': True},
    'split':      {'title': 'Dividir PDF',    'desc': 'Separa cada página en un archivo',             'icon': '⊗', 'accept': '.pdf',                  'multiple': False},
    'compress':   {'title': 'Comprimir PDF',  'desc': 'Reduce el tamaño del archivo PDF',             'icon': '◎', 'accept': '.pdf',                  'multiple': False},
    'pdf-to-word':{'title': 'PDF a Word',     'desc': 'Convierte PDF a documento Word editable',      'icon': '→', 'accept': '.pdf',                  'multiple': False},
    'pdf-to-jpg': {'title': 'PDF a JPG',      'desc': 'Convierte páginas PDF a imágenes JPG',         'icon': '◉', 'accept': '.pdf',                  'multiple': False},
    'word-to-pdf':{'title': 'Word a PDF',     'desc': 'Convierte documentos Word a PDF',              'icon': '←', 'accept': '.docx',                 'multiple': False},
    'jpg-to-pdf': {'title': 'JPG a PDF',      'desc': 'Convierte imágenes JPG/PNG a PDF',             'icon': '◈', 'accept': '.jpg,.jpeg,.png,.webp', 'multiple': True},
    'protect':    {'title': 'Proteger PDF',   'desc': 'Añade contraseña a tu PDF',                    'icon': '⊘', 'accept': '.pdf',                  'multiple': False},
    'rotate':     {'title': 'Rotar PDF',      'desc': 'Rota las páginas de tu PDF',                   'icon': '↻', 'accept': '.pdf',                  'multiple': False},
}

@app.route('/')
def index():
    return render_template('index.html', tools=TOOLS)

@app.route('/tool/<n>')
def tool(n):
    if n not in TOOLS:
        return render_template('404.html'), 404
    return render_template('tool.html', tool_name=n, tool=TOOLS[n])

# MERGE
@app.route('/api/merge', methods=['POST'])
def api_merge():
    files = request.files.getlist('files')
    if len(files) < 2:
        return jsonify({'error': 'Se necesitan al menos 2 PDFs'}), 400
    writer = PdfWriter()
    for f in files:
        if not allowed(f.filename, ALLOWED_PDF):
            return jsonify({'error': f'{f.filename} no es PDF'}), 400
        for page in PdfReader(f).pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')

# SPLIT
@app.route('/api/split', methods=['POST'])
def api_split():
    f = request.files.get('file')
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    reader = PdfReader(f)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(reader.pages):
            w = PdfWriter()
            w.add_page(page)
            pb = io.BytesIO()
            w.write(pb)
            zf.writestr(f'page_{i+1}.pdf', pb.getvalue())
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype='application/zip', as_attachment=True, download_name='pages.zip')

# COMPRESS
@app.route('/api/compress', methods=['POST'])
def api_compress():
    f = request.files.get('file')
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    reader = PdfReader(f)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    for page in writer.pages:
        page.compress_content_streams()
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='compressed.pdf')

# PDF TO WORD
@app.route('/api/pdf-to-word', methods=['POST'])
def api_pdf_to_word():
    f = request.files.get('file')
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    reader = PdfReader(f)
    doc = Document()
    doc.add_heading('Documento convertido', 0)
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ''
        doc.add_heading(f'Página {i+1}', level=2)
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='converted.docx')

# PDF TO JPG
@app.route('/api/pdf-to-jpg', methods=['POST'])
def api_pdf_to_jpg():
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        return jsonify({'error': 'Instala pdf2image y poppler-utils'}), 500
    f = request.files.get('file')
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    images = convert_from_bytes(f.read(), dpi=150)
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, img in enumerate(images):
            ib = io.BytesIO()
            img.save(ib, 'JPEG', quality=85)
            zf.writestr(f'page_{i+1}.jpg', ib.getvalue())
    zip_buf.seek(0)
    return send_file(zip_buf, mimetype='application/zip', as_attachment=True, download_name='images.zip')

# WORD TO PDF
@app.route('/api/word-to-pdf', methods=['POST'])
def api_word_to_pdf():
    f = request.files.get('file')
    if not f or not allowed(f.filename, ALLOWED_DOCX):
        return jsonify({'error': 'Archivo .docx requerido'}), 400
    try:
        import subprocess
        tmp_in = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        f.save(tmp_in.name); tmp_in.close()
        tmp_dir = tempfile.mkdtemp()
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, tmp_in.name],
                       capture_output=True, timeout=30)
        os.unlink(tmp_in.name)
        pdf_files = [x for x in os.listdir(tmp_dir) if x.endswith('.pdf')]
        if not pdf_files:
            return jsonify({'error': 'Conversión fallida. ¿LibreOffice instalado?'}), 500
        return send_file(os.path.join(tmp_dir, pdf_files[0]), mimetype='application/pdf',
                         as_attachment=True, download_name='converted.pdf')
    except FileNotFoundError:
        return jsonify({'error': 'LibreOffice no encontrado: sudo apt install libreoffice'}), 500

# JPG TO PDF
@app.route('/api/jpg-to-pdf', methods=['POST'])
def api_jpg_to_pdf():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'Se necesita al menos una imagen'}), 400
    images_bytes = []
    for f in files:
        if not allowed(f.filename, ALLOWED_IMG):
            return jsonify({'error': f'{f.filename} no es imagen válida'}), 400
        img = Image.open(f).convert('RGB')
        buf = io.BytesIO()
        img.save(buf, 'JPEG', quality=90)
        images_bytes.append(buf.getvalue())
    pdf_bytes = img2pdf.convert(images_bytes)
    buf = io.BytesIO(pdf_bytes); buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='images.pdf')

# PROTECT
@app.route('/api/protect', methods=['POST'])
def api_protect():
    f = request.files.get('file')
    password = request.form.get('password', '')
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    if not password:
        return jsonify({'error': 'Contraseña requerida'}), 400
    reader = PdfReader(f)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    buf = io.BytesIO(); writer.write(buf); buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='protected.pdf')

# ROTATE
@app.route('/api/rotate', methods=['POST'])
def api_rotate():
    f = request.files.get('file')
    degrees = int(request.form.get('degrees', 90))
    if not f or not allowed(f.filename, ALLOWED_PDF):
        return jsonify({'error': 'PDF requerido'}), 400
    if degrees not in [90, 180, 270]:
        return jsonify({'error': 'Grados inválidos'}), 400
    reader = PdfReader(f)
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    buf = io.BytesIO(); writer.write(buf); buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='rotated.pdf')

if __name__ == '__main__':
    app.run(debug=True, port=5000)
